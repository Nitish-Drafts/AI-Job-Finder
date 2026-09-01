import os
import json
import logging
from typing import Dict, List, Any, Optional
import httpx
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from app.core.config import settings

logger = logging.getLogger("optimization_service")


class OptimizationService:
    """
    Handles Resume Optimization, Cover Letter Generation, Interview Q&A Generation,
    and Document Exports (PDF/DOCX).
    """

    async def optimize_resume(self, resume_text: str, job_title: str, job_description: str) -> Dict[str, Any]:
        """
        Generates ATS-optimized summaries, targeted skill lists, and rewritten project descriptions.
        """
        if settings.OPENAI_API_KEY:
            try:
                prompt = f"""
                You are an expert ATS optimizer and resume writer.
                Optimize the following resume text for the target job: "{job_title}".
                
                Target Job Description:
                {job_description}
                
                Current Resume Text:
                {resume_text}
                
                Return ONLY a JSON object with this structure:
                {{
                    "optimized_summary": "A professional summary (3-4 sentences) optimized with keywords from the JD.",
                    "optimized_skills": ["Skill1", "Skill2", "Skill3"],
                    "improved_projects": [
                        {{
                            "original_title": "Project Title",
                            "improved_description": "A metrics-driven, impact-oriented description incorporating keywords from the JD."
                        }}
                    ],
                    "ats_tips": ["Tip1", "Tip2"]
                }}
                """
                async with httpx.AsyncClient(timeout=30.0) as client:
                    response = await client.post(
                        "https://api.openai.com/v1/chat/completions",
                        headers={"Authorization": f"Bearer {settings.OPENAI_API_KEY}"},
                        json={
                            "model": "gpt-4o-mini",
                            "messages": [{"role": "user", "content": prompt}],
                            "response_format": {"type": "json_object"},
                            "temperature": 0.3
                        }
                    )
                    if response.status_code == 200:
                        return json.loads(response.json()["choices"][0]["message"]["content"])
            except Exception as e:
                logger.error(f"OpenAI resume optimization failed: {e}. Using local template.")

        # Local Fallback Optimization
        # Extract skills from job description
        from app.services.ai_service import TECH_KEYWORDS
        detected_jd_skills = []
        desc_lower = job_description.lower()
        for kw in TECH_KEYWORDS[:100]:
            if kw in desc_lower:
                detected_jd_skills.append(kw.capitalize())
                
        optimized_skills = list(set(detected_jd_skills[:12]))
        
        return {
            "optimized_summary": f"Results-driven Software Engineer with extensive experience developing high-performance services. Proven expertise in building web interfaces, maintaining secure cloud architectures, and utilizing modern stacks including {', '.join(optimized_skills[:4])}. Adept at collaborating with cross-functional teams to deploy features in alignment with target goals for {job_title}.",
            "optimized_skills": optimized_skills or ["Software Engineering", "Python", "React", "Docker"],
            "improved_projects": [
                {
                    "original_title": "Core Platform Development",
                    "improved_description": f"Architected and deployed microservices backend matching target specs for {job_title}, incorporating containerization with Docker and automated workflows via CI/CD, resulting in a 25% reduction in production hotfixes."
                },
                {
                    "original_title": "Interactive Client Dashboard",
                    "improved_description": "Engineered user-facing dashboard using React and TypeScript, optimizing payload rendering and REST interface connections to reduce client-side loading times by 40%."
                }
            ],
            "ats_tips": [
                "Incorporate metric-driven bullet points focusing on percentages and hours saved.",
                "Ensure your skill names match the exact spelling found in the job description."
            ]
        }

    async def generate_cover_letter(
        self, resume_text: str, job_title: str, company_name: str, job_description: str
    ) -> str:
        """
        Generates a professionally drafted cover letter.
        """
        if settings.OPENAI_API_KEY:
            try:
                prompt = f"""
                Write a professional cover letter for a candidate applying for the "{job_title}" role at "{company_name}".
                Align the candidate's experience and skills from the resume with the requirements of the job description.
                Keep it to 3-4 paragraphs. Make it look formal and ready to sign.
                
                Target Job Description:
                {job_description}
                
                Candidate Resume Text:
                {resume_text}
                
                Write only the text of the cover letter.
                """
                async with httpx.AsyncClient(timeout=30.0) as client:
                    response = await client.post(
                        "https://api.openai.com/v1/chat/completions",
                        headers={"Authorization": f"Bearer {settings.OPENAI_API_KEY}"},
                        json={
                            "model": "gpt-4o-mini",
                            "messages": [{"role": "user", "content": prompt}],
                            "temperature": 0.4
                        }
                    )
                    if response.status_code == 200:
                        return response.json()["choices"][0]["message"]["content"].strip()
            except Exception as e:
                logger.error(f"OpenAI cover letter generation failed: {e}. Using local template.")

        # Local Template Cover Letter
        from app.services.ai_service import TECH_KEYWORDS
        detected = [kw.capitalize() for kw in TECH_KEYWORDS[:50] if kw in job_description.lower()]
        skills_str = ", ".join(detected[:4]) if detected else "software design, agile operations, and container deployments"

        cover_letter = f"""[Your Name]
[Your Address]
[Your Phone Number] | [Your Email]

{datetime.utcnow().strftime('%B %d, %Y')}

Hiring Team
{company_name}
[Company Address]

Subject: Application for {job_title} position

Dear Hiring Team,

I am writing to express my enthusiastic interest in the {job_title} position currently open at {company_name}. With my background in software engineering and hands-on experience developing responsive applications, I am confident that my technical skills and passion for building scalable solutions will align perfectly with your team's objectives.

Throughout my career, I have specialized in technology stacks including {skills_str}. In my previous projects, I successfully built secure service infrastructures, optimized database performance, and refined deployment pipelines, which improved release cycles and client satisfaction. I admire {company_name}'s dedication to innovation and quality, and I am eager to apply my skills to help solve your engineering challenges.

I would welcome the opportunity to discuss how my qualifications, collaborative mindset, and technical background make me a strong addition to your team. Thank you for your time and consideration.

Sincerely,

[Your Name]"""
        return cover_letter

    async def generate_interview_questions(
        self, job_title: str, company_name: str, job_description: str
    ) -> List[Dict[str, Any]]:
        """
        Generates technical, behavioral, coding, SQL, Python, and ML questions.
        """
        if settings.OPENAI_API_KEY:
            try:
                prompt = f"""
                Generate a list of 6 highly targeted interview questions for a candidate interviewing for the "{job_title}" role at "{company_name}".
                
                Target Job Description:
                {job_description}
                
                Include:
                - 1 Technical question
                - 1 Behavioral question
                - 1 Coding question
                - 1 SQL question
                - 1 Python question
                - 1 Machine Learning / System Design question
                
                Return ONLY a JSON array containing objects matching this structure:
                [
                    {{
                        "question_type": "technical | behavioral | coding | sql | python | ml",
                        "question": "The question text",
                        "expected_answer": "Summarized bullet points of what a strong answer should cover.",
                        "difficulty": "easy | medium | hard"
                    }}
                ]
                """
                async with httpx.AsyncClient(timeout=30.0) as client:
                    response = await client.post(
                        "https://api.openai.com/v1/chat/completions",
                        headers={"Authorization": f"Bearer {settings.OPENAI_API_KEY}"},
                        json={
                            "model": "gpt-4o-mini",
                            "messages": [{"role": "user", "content": prompt}],
                            "response_format": {"type": "json_object" if settings.OPENAI_API_KEY else "text"}, # Wait, json array is expected
                            "temperature": 0.4
                        }
                    )
                    if response.status_code == 200:
                        # Some models need JSON object wrapping
                        res_content = response.json()["choices"][0]["message"]["content"]
                        data = json.loads(res_content)
                        if isinstance(data, list):
                            return data
                        elif isinstance(data, dict) and "questions" in data:
                            return data["questions"]
                        elif isinstance(data, dict):
                            return list(data.values())[0]
            except Exception as e:
                logger.error(f"OpenAI interview generation failed: {e}. Using local template.")

        # Local Q&A Fallback
        return [
            {
                "question_type": "technical",
                "question": f"Explain the architectural challenges when scaling a platform like {company_name}'s core application.",
                "expected_answer": "Candidates should address database partitioning/sharding, caching layers (Redis), rate-limiting, and load balancing high traffic volumes.",
                "difficulty": "hard"
            },
            {
                "question_type": "behavioral",
                "question": "Describe a time you faced a critical bug in production right before a release. How did you coordinate the hotfix?",
                "expected_answer": "Look for strong communication, calm isolation of the issue using logs, rollback strategies, and implementing regression tests post-mortem.",
                "difficulty": "medium"
            },
            {
                "question_type": "coding",
                "question": "Given a list of job search coordinates, find the top k closest jobs using Euclidean distance. Write a helper function in O(N log k) time complexity.",
                "expected_answer": "Should utilize a min-heap or max-heap structure to keep track of the closest jobs efficiently without sorting the entire array.",
                "difficulty": "medium"
            },
            {
                "question_type": "sql",
                "question": "Write a query to find the average ATS match score of applications grouped by month for the past year, filtering out categories with less than 5 candidates.",
                "expected_answer": "SELECT DATE_TRUNC('month', applied_date) as m, AVG(score) FROM apps GROUP BY m HAVING COUNT(id) >= 5;",
                "difficulty": "medium"
            },
            {
                "question_type": "python",
                "question": "How do you handle thread-safety and race conditions when scraping multiple job websites concurrently in Python?",
                "expected_answer": "Explain asyncio event loops, standard semaphore controls (`asyncio.Semaphore`), thread locks for synchronous database writing, and request throttling.",
                "difficulty": "hard"
            },
            {
                "question_type": "ml",
                "question": "How would you design a Job Recommendation Engine? What metrics would you evaluate?",
                "expected_answer": "Propose collaborative filtering, content-based matching using TF-IDF/Sentence Embeddings, and evaluate using Precision@K, Recall@K, and NDCG.",
                "difficulty": "hard"
            }
        ]

    def export_docx(self, content_text: str, file_path: str) -> None:
        """
        Saves text content as a DOCX document.
        """
        doc = Document()
        
        # Split text into paragraphs and write
        paragraphs = content_text.split("\n")
        for para in paragraphs:
            para = para.strip()
            if para:
                doc.add_paragraph(para)
            else:
                doc.add_paragraph()  # spacer
                
        doc.save(file_path)

    def export_pdf(self, content_text: str, file_path: str) -> None:
        """
        Saves text content as a styled PDF document using ReportLab.
        """
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create a custom body style
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=8
        )
        
        story = []
        paragraphs = content_text.split("\n")
        for para in paragraphs:
            para = para.strip()
            if para:
                story.append(Paragraph(para, body_style))
            else:
                story.append(Spacer(1, 10))
                
        doc.build(story)


# Global instance of service
optimization_service = OptimizationService()
from datetime import datetime
